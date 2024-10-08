import streamlit as st
from docx import Document
import tempfile

def process_docx(doc, cursive_prefix, normal_prefix):
    processed_lines = []
    for para in doc.paragraphs:
        lines = para.text.split('\n')
        for line in lines:
            stripped_line = line.strip()
            if stripped_line:
                is_italic = False
                for run in para.runs:
                    if run.text in line and run.font.italic:
                        is_italic = True
                        break
                if is_italic:
                    processed_lines.append(cursive_prefix + stripped_line)
                else:
                    processed_lines.append(normal_prefix + stripped_line)
    return '\n'.join(processed_lines)

def main():
    st.title("Interview Helper")

    uploaded_file = st.file_uploader("Bitte eine Word-Datei (docx) hochladen", type="docx")

    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(uploaded_file.getvalue())
            temp_file.close()

            doc = Document(temp_file.name)
            
            st.write("Original-Text:")
            st.text("\n".join([para.text for para in doc.paragraphs]))

            #cursive_prefix = st.text_input("Enter prefix for cursive text:", "Interviewer: ")
            cursive_prefix = "Profil: "
            normal_prefix = st.text_input("Name der Person:", "Interviewter") + ": "

            if st.button("Anpassen"):
                processed_text = process_docx(doc, cursive_prefix, normal_prefix)
                
                st.write("Angepasster Text:")
                st.code(processed_text)

if __name__ == "__main__":
    main()

